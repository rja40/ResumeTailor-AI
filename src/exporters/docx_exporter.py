"""Export Markdown / plain text to DOCX bytes."""
from __future__ import annotations

import io
import re

from docx import Document
from docx.shared import Pt


def markdown_to_docx_bytes(md: str) -> bytes:
    """Render a subset of Markdown (#/##/###, bullets, bold) into DOCX."""
    doc = Document()
    _set_default_font(doc)

    for raw_line in md.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            content = line.lstrip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _write_runs_with_bold(p, content)
        else:
            p = doc.add_paragraph()
            _write_runs_with_bold(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def text_to_docx_bytes(text: str) -> bytes:
    """Plain text → DOCX (one paragraph per line)."""
    doc = Document()
    _set_default_font(doc)
    for line in text.splitlines():
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _write_runs_with_bold(paragraph, text: str) -> None:
    cursor = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > cursor:
            paragraph.add_run(text[cursor : m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        cursor = m.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])
