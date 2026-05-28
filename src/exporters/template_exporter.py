"""Rewrite a DOCX in place to preserve the original design.

We walk every paragraph (including nested tables, which are how most
two-column resumes are laid out), let the caller decide which paragraphs
to replace by index, and inject the new text while keeping the first run's
font, size, color, and bold/italic state. Headings, names, dates, and
company-name paragraphs are never touched.
"""
from __future__ import annotations

import io
from typing import Iterator

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


_MC_FALLBACK = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"


def _is_in_mc_fallback(elem) -> bool:
    parent = elem.getparent()
    while parent is not None:
        if parent.tag == _MC_FALLBACK:
            return True
        parent = parent.getparent()
    return False


def iter_all_paragraphs(doc) -> Iterator[Paragraph]:
    """Yield every <w:p> in the document — body, tables, AND text boxes.

    Designer-style resumes build column layouts with floating text boxes
    (<w:txbxContent>); python-docx's `doc.paragraphs` skips those. Iterating
    the underlying lxml tree for all <w:p> elements catches them uniformly.

    Skips paragraphs inside <mc:Fallback> branches so we don't see each text-
    box paragraph twice (Word stores a modern + legacy rendering).
    """
    for p_elem in doc.element.body.iter(qn("w:p")):
        if _is_in_mc_fallback(p_elem):
            continue
        yield Paragraph(p_elem, doc)


def collect_rewritable_paragraphs(source_bytes: bytes) -> list[dict]:
    """Return [{idx, text, kind}] for paragraphs that look like rewritable content.

    A paragraph is rewritable when it's either:
      - styled as a bullet/list item (any style name containing "list")
      - or normal-style prose of >= 50 characters (summaries / role bullets that
        weren't tagged with a list style)

    Headings, names, contact lines, role titles, dates, and company names are
    deliberately excluded.
    """
    doc = Document(io.BytesIO(source_bytes))
    out: list[dict] = []
    for idx, p in enumerate(iter_all_paragraphs(doc)):
        text = p.text.strip()
        if not text:
            continue
        style_name = (p.style.name if p.style else "").lower()
        if "heading" in style_name or "title" in style_name:
            continue
        if "list" in style_name:
            out.append({"idx": idx, "text": text, "kind": "bullet"})
        elif len(text) >= 50:
            out.append({"idx": idx, "text": text, "kind": "prose"})
    return out


def apply_rewrites(source_bytes: bytes, rewrites: dict[int, str]) -> bytes:
    """Replace paragraph text by global index. Returns new DOCX bytes.

    Formatting is preserved by writing the new text into the paragraph's first
    run (which carries the font/size/color) and clearing subsequent runs.
    """
    doc = Document(io.BytesIO(source_bytes))
    paragraphs = list(iter_all_paragraphs(doc))
    for raw_idx, new_text in rewrites.items():
        idx = int(raw_idx)
        if not (0 <= idx < len(paragraphs)) or not isinstance(new_text, str):
            continue
        _replace_keep_style(paragraphs[idx], new_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render_preview_markdown(source_bytes: bytes) -> str:
    """Render the (post-rewrite) DOCX as a simple Markdown preview for the UI."""
    doc = Document(io.BytesIO(source_bytes))
    lines: list[str] = []
    for p in iter_all_paragraphs(doc):
        text = p.text.strip()
        if not text:
            lines.append("")
            continue
        style_name = (p.style.name if p.style else "").lower()
        if "heading 1" in style_name or "title" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif "list" in style_name:
            lines.append(f"- {text}")
        else:
            lines.append(text)
    return "\n".join(lines).strip()


def _replace_keep_style(paragraph: Paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first = paragraph.runs[0]
    first.text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
